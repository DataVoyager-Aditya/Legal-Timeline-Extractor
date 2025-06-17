
import logging
from pathlib import Path
from docx import Document
import zipfile
from typing import Dict, Any

class TextProcessor:
    """Process text files (.txt, .docx, .rtf)"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text(self, file_path: str) -> str:
        """Extract text from various text file formats"""
        try:
            path = Path(file_path)

            if path.suffix.lower() == '.txt':
                return self._extract_from_txt(file_path)
            elif path.suffix.lower() == '.docx':
                return self._extract_from_docx(file_path)
            elif path.suffix.lower() == '.rtf':
                return self._extract_from_rtf(file_path)
            else:
                raise ValueError(f"Unsupported text format: {path.suffix}")

        except Exception as e:
            self.logger.error(f"Error extracting text: {e}")
            return ""

    def _extract_from_txt(self, txt_path: str) -> str:
        """Extract text from plain text file"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # Final attempt with error handling
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_from_docx(self, docx_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(docx_path)
            text_content = []

            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_content.append(" | ".join(row_text))

            return "\n".join(text_content)

        except Exception as e:
            self.logger.error(f"Error extracting DOCX: {e}")
            return ""

    def _extract_from_rtf(self, rtf_path: str) -> str:
        """Extract text from RTF file"""
        try:
            # Simple RTF text extraction
            with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            # Remove RTF formatting codes (basic cleanup)
            import re
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            # Remove RTF control symbols
            content = re.sub(r'\\[^a-z]', '', content)
            # Remove braces
            content = re.sub(r'[{}]', '', content)

            return content

        except Exception as e:
            self.logger.error(f"Error extracting RTF: {e}")
            return ""
