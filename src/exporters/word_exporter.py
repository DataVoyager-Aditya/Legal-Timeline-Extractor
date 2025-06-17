
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from io import BytesIO
import logging
from typing import List, Dict, Any

class WordExporter:
    """
    Professional Word document exporter for legal timeline reports.
    Creates formatted documents suitable for legal proceedings.
    """

    def __init__(self):
        """Initialize Word exporter"""
        self.logger = logging.getLogger(__name__)

    def create_timeline_document(self, timeline_events: List[Dict[str, Any]], metadata: Dict[str, Any]) -> BytesIO:
        """
        Create comprehensive Word document with timeline analysis.

        Args:
            timeline_events: List of timeline events
            metadata: Report metadata

        Returns:
            BytesIO buffer containing Word document data
        """
        buffer = BytesIO()

        try:
            # Create document
            doc = Document()

            # Set up styles
            self._setup_document_styles(doc)

            # Create content
            self._create_cover_page(doc, metadata)
            self._add_page_break(doc)
            self._create_executive_summary(doc, timeline_events, metadata)
            self._add_page_break(doc)
            self._create_timeline_section(doc, timeline_events)
            self._add_page_break(doc)
            self._create_evidence_section(doc, timeline_events)

            # Save to buffer
            doc.save(buffer)
            buffer.seek(0)

            self.logger.info("Word document generated successfully")
            return buffer

        except Exception as e:
            self.logger.error(f"Error creating Word document: {e}")
            raise

    def _setup_document_styles(self, doc: Document):
        """Setup document styles"""
        styles = doc.styles

        # Title style
        title_style = styles['Title']
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.color.rgb = RGBColor(30, 60, 114)  # Legal blue

        # Heading styles
        heading1_style = styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(16)
        heading1_font.color.rgb = RGBColor(42, 82, 152)

        # Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(11)

    def _create_cover_page(self, doc: Document, metadata: Dict[str, Any]):
        """Create document cover page"""
        # Title
        title = doc.add_heading('LEGAL TIMELINE ANALYSIS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph('AI-Powered Document Analysis')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.italic = True

        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Report information table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Populate table
        info_data = [
            ("Generated:", datetime.now().strftime('%B %d, %Y at %I:%M %p')),
            ("Total Events:", str(metadata.get('total_events', 0))),
            ("Source Documents:", str(len(metadata.get('source_files', [])))),
            ("Analysis Method:", metadata.get('extraction_method', 'Legal-BERT AI')),
            ("System Version:", metadata.get('system_version', 'AI Legal Timeline Builder Pro'))
        ]

        for i, (label, value) in enumerate(info_data):
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value

            # Bold the labels
            row_cells[0].paragraphs[0].runs[0].font.bold = True

        # Source documents section
        if metadata.get('source_files'):
            doc.add_paragraph()
            doc.add_heading('Source Documents:', level=2)

            for i, filename in enumerate(metadata['source_files'], 1):
                p = doc.add_paragraph(f"{i}. {filename}")
                p.style = 'List Number'

        # Disclaimer
        doc.add_paragraph()
        doc.add_paragraph()
        disclaimer_p = doc.add_paragraph()
        disclaimer_run = disclaimer_p.add_run('DISCLAIMER: ')
        disclaimer_run.font.bold = True
        disclaimer_p.add_run(
            'This timeline analysis was generated using artificial intelligence '
            'and should be reviewed by qualified legal professionals. All source '
            'documents should be independently verified for accuracy and completeness.'
        )
        disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _create_executive_summary(self, doc: Document, timeline_events: List[Dict[str, Any]], metadata: Dict[str, Any]):
        """Create executive summary section"""
        doc.add_heading('EXECUTIVE SUMMARY', level=1)

        # Summary statistics
        total_events = len(timeline_events)
        high_confidence = len([e for e in timeline_events if e.get('confidence', 0) > 0.8])
        date_range = self._get_date_range(timeline_events)

        summary_text = (
            f"This report presents a chronological analysis of {total_events} events extracted from "
            f"{len(metadata.get('source_files', []))} legal documents using advanced AI technology. "
            f"The analysis identified {high_confidence} high-confidence events spanning from {date_range}.\n\n"

            "The timeline extraction process utilized Legal-BERT, a specialized natural language processing "
            "model trained on legal text, combined with pattern recognition algorithms optimized for "
            "legal document analysis. Each event has been assigned a confidence score based on the "
            "clarity of the source text and the reliability of the extraction method."
        )

        p = doc.add_paragraph(summary_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Key findings
        doc.add_heading('Key Findings:', level=2)

        # Statistics table
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'

        avg_confidence = sum(e.get('confidence', 0) for e in timeline_events) / total_events if total_events > 0 else 0

        stats_data = [
            ("Total Events", str(total_events)),
            ("High Confidence Events (>80%)", str(high_confidence)),
            ("Average Confidence", f"{avg_confidence:.1%}"),
            ("Date Range", date_range),
            ("Source Documents", str(len(metadata.get('source_files', []))))
        ]

        for i, (label, value) in enumerate(stats_data):
            row_cells = stats_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            row_cells[0].paragraphs[0].runs[0].font.bold = True

    def _create_timeline_section(self, doc: Document, timeline_events: List[Dict[str, Any]]):
        """Create detailed timeline section"""
        doc.add_heading('CHRONOLOGICAL TIMELINE', level=1)

        for i, event in enumerate(timeline_events, 1):
            # Event header
            event_heading = f"{i}. {event.get('date', 'Unknown Date')} - {event.get('event', 'Unknown Event')}"
            doc.add_heading(event_heading, level=2)

            # Event details table
            details_table = doc.add_table(rows=4, cols=2)
            details_table.style = 'Table Grid'

            # Populate details
            entities = ', '.join(event.get('entities', [])) or 'None identified'
            confidence = f"{event.get('confidence', 0):.1%}"
            source_file = event.get('metadata', {}).get('filename', 'Unknown Source')
            description = event.get('text', 'No description available')[:200] + ('...' if len(event.get('text', '')) > 200 else '')

            details_data = [
                ("Parties/Entities:", entities),
                ("Confidence Score:", confidence),
                ("Source Document:", source_file),
                ("Description:", description)
            ]

            for j, (label, value) in enumerate(details_data):
                row_cells = details_table.rows[j].cells
                row_cells[0].text = label
                row_cells[1].text = value
                row_cells[0].paragraphs[0].runs[0].font.bold = True

            doc.add_paragraph()  # Add spacing

    def _create_evidence_section(self, doc: Document, timeline_events: List[Dict[str, Any]]):
        """Create evidence appendix section"""
        doc.add_heading('EVIDENCE APPENDIX', level=1)

        # Group events by source
        sources = {}
        for event in timeline_events:
            source = event.get('metadata', {}).get('filename', 'Unknown Source')
            if source not in sources:
                sources[source] = []
            sources[source].append(event)

        for source_file, events in sources.items():
            doc.add_heading(f"Source Document: {source_file}", level=2)

            p = doc.add_paragraph(f"Events extracted from this document: {len(events)}")

            # Create evidence table
            evidence_table = doc.add_table(rows=len(events) + 1, cols=3)
            evidence_table.style = 'Table Grid'

            # Headers
            header_cells = evidence_table.rows[0].cells
            header_cells[0].text = "Date"
            header_cells[1].text = "Event"
            header_cells[2].text = "Confidence"

            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True

            # Data rows
            for j, event in enumerate(events, 1):
                row_cells = evidence_table.rows[j].cells
                row_cells[0].text = event.get('date', 'Unknown')
                row_cells[1].text = event.get('event', 'Unknown')[:50] + ('...' if len(event.get('event', '')) > 50 else '')
                row_cells[2].text = f"{event.get('confidence', 0):.1%}"

            doc.add_paragraph()  # Add spacing

    def _add_page_break(self, doc: Document):
        """Add page break to document"""
        doc.add_page_break()

    def _get_date_range(self, timeline_events: List[Dict[str, Any]]) -> str:
        """Get date range from timeline events"""
        dates = [e.get('date', '') for e in timeline_events if e.get('date') and e.get('date') != 'Unknown']
        if dates:
            return f"{min(dates)} to {max(dates)}"
        return "Date range unavailable"
